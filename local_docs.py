# local_docs.py
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
import os
from datetime import datetime

def create_local_doc(title, content, output_dir="generated_documents"):
    """
    Создаёт профессиональный медицинский протокол в формате .docx
    с поддержкой кириллицы и красивым форматированием.
    """
    # Создаём директорию, если её нет
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    # Создаём документ
    doc = Document()
    
    # Устанавливаем альбомную ориентацию для лучшего отображения
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    
    # Заголовок
    header = doc.add_heading(title, level=0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header.runs[0]
    header_run.font.size = Pt(16)
    header_run.font.bold = True
    
    # Подзаголовок — дата
    date_para = doc.add_paragraph(f"Дата: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    # Разбиваем контент на секции
    lines = content.strip().split('\n')
    for line in lines:
        if line.startswith("**") and line.endswith("**"):
            # Это заголовок секции
            heading = doc.add_heading(line.strip("* "), level=2)
            heading_run = heading.runs[0]
            heading_run.font.bold = True
        elif line.startswith("- "):
            # Это пункт списка
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.strip():
            # Обычный текст
            p = doc.add_paragraph(line)
            p.style.font.size = Pt(12)
    
    # Добавляем подпись
    doc.add_page_break()
    signature = doc.add_paragraph()
    signature.add_run("Врач: ____________________ /ФИО/\n")
    signature.add_run(f"Дата: {datetime.now().strftime('%d.%m.%Y')}")
    signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Генерируем имя файла
    safe_title = "".join(c for c in title if c.isalnum() or c in " -_").rstrip()
    filename = f"{safe_title}.docx"
    filepath = os.path.join(output_dir, filename)
    
    # Сохраняем
    doc.save(filepath)
    
    return filepath, f"✅ Протокол сохранён: {filepath}"